import streamlit as st
import pandas as pd
import json
import os
import io
from datetime import datetime
from sqlalchemy import create_engine, Column, Integer, String, Float, Boolean, DateTime, JSON, ForeignKey
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import plotly.express as px

# ---------- 数据库模型 ----------
Base = declarative_base()

class Project(Base):
    __tablename__ = 'projects'
    id = Column(Integer, primary_key=True)
    name = Column(String(200), nullable=False)
    route = Column(String(100))
    pile_no = Column(String(50))
    side = Column(String(10))
    type = Column(String(10))
    zone = Column(String(10))
    design_height = Column(Float)
    current_height = Column(Float)
    design_length = Column(Float)
    current_length = Column(Float)
    design_slope = Column(String(20))
    actual_slope = Column(String(20))
    grade_count = Column(Integer)
    platform_count = Column(Integer)
    traffic_volume = Column(Integer)
    facilities = Column(String(200))
    monitor_emergency = Column(String(30))
    operating_years = Column(Integer)
    rainfall_3d = Column(Float)
    rainfall_7d = Column(Float)
    fs_input = Column(Float, nullable=True)
    created_at = Column(DateTime, default=datetime.now)
    updated_at = Column(DateTime, default=datetime.now, onupdate=datetime.now)

class SurveyDefect(Base):
    __tablename__ = 'survey_defects'
    id = Column(Integer, primary_key=True)
    project_id = Column(Integer, ForeignKey('projects.id'), nullable=True)
    category = Column(String(1))
    code = Column(String(20))
    checked = Column(Boolean, default=False)
    description = Column(String(200))
    level2 = Column(String(50))
    level3 = Column(String(50))
    level4 = Column(String(50))
    base_deduction = Column(Float)
    zone_factor = Column(Float, default=1.0)

class WeightConfig(Base):
    __tablename__ = 'weight_configs'
    id = Column(Integer, primary_key=True)
    project_id = Column(Integer, ForeignKey('projects.id'), nullable=True)
    weight_D = Column(Float, default=0.20)
    weight_C = Column(Float, default=0.25)
    weight_E = Column(Float, default=0.35)
    weight_O = Column(Float, default=0.20)
    weight_a1 = Column(Float, default=0.40)
    weight_a2 = Column(Float, default=0.30)
    weight_a3 = Column(Float, default=0.20)
    weight_a4 = Column(Float, default=0.10)
    zone_factors = Column(JSON)
    coupling_deductions = Column(JSON)
    is_global = Column(Boolean, default=True)

class SystemConfig(Base):
    __tablename__ = 'system_configs'
    id = Column(Integer, primary_key=True)
    config_type = Column(String(50), unique=True)
    config_data = Column(JSON)
    updated_at = Column(DateTime, default=datetime.now, onupdate=datetime.now)

class Result(Base):
    __tablename__ = 'results'
    id = Column(Integer, primary_key=True)
    project_id = Column(Integer, ForeignKey('projects.id'))
    score_D = Column(Float)
    score_C = Column(Float)
    score_E = Column(Float)
    score_O = Column(Float)
    S_base = Column(Float)
    coupling_total = Column(Float)
    S_base_corrected = Column(Float)
    alpha1 = Column(Float)
    alpha2 = Column(Float)
    alpha3 = Column(Float)
    alpha4 = Column(Float)
    alpha = Column(Float)
    S_final = Column(Float)
    initial_grade = Column(String(10))
    final_grade = Column(String(10))
    response_measures = Column(String(500))
    created_at = Column(DateTime, default=datetime.now)

# engine = create_engine('sqlite:///slope_risk.db', echo=False)
import os
import streamlit as st

# 优先从 Streamlit Secrets 读取，其次从环境变量读取
DATABASE_URL = st.secrets.get("DATABASE_URL", os.environ.get("DATABASE_URL", "postgresql://postgres:qwedcxzas8922LJY@aws-1-ap-southeast-1.supabase.co:5432/postgres"))
engine = create_engine(DATABASE_URL, echo=False)
Base.metadata.create_all(engine)
Session = sessionmaker(bind=engine)

# ---------- 初始化系统配置 ----------
def init_system_configs():
    session = Session()
    if not session.query(SystemConfig).filter_by(config_type='grade_thresholds').first():
        default = {'thresholds': [
            {'min': 80, 'max': 100, 'level': '低风险', 'color': '蓝色'},
            {'min': 60, 'max': 79, 'level': '中风险', 'color': '黄色'},
            {'min': 40, 'max': 59, 'level': '高风险', 'color': '橙色'},
            {'min': 0,  'max': 39, 'level': '极高风险', 'color': '红色'}
        ]}
        session.add(SystemConfig(config_type='grade_thresholds', config_data=default))
    if not session.query(SystemConfig).filter_by(config_type='response_measures').first():
        default = {
            '低风险': '常规巡检，每年一次定期评估。',
            '中风险': '加强日常巡检（每月专项巡查1次），半年内制定处治方案，必要时设临时监测。',
            '高风险': '限速60km/h，设置临时排水，3个月内完成加固设计，启动自动化监测。',
            '极高风险': '立即封闭车道或路段，24小时监控，启动应急抢险，组织专家会商。'
        }
        session.add(SystemConfig(config_type='response_measures', config_data=default))
    if not session.query(SystemConfig).filter_by(config_type='safety_factors').first():
        default = {'正常工况': 1.25, '暴雨工况': 1.15, '地震工况': 1.10}
        session.add(SystemConfig(config_type='safety_factors', config_data=default))
    if not session.query(SystemConfig).filter_by(config_type='coupling_rules').first():
        default = {'rules': [
            {'code': 'PC1', 'desc': '排水失效+强降雨', 'deduction': -8, 'triggers': ['D4.2_1','C4.1_4','O1.6_7'], 'extra': 'rainfall_7d>100'},
            {'code': 'PC2', 'desc': '顺层+持续降雨', 'deduction': -6, 'triggers': ['D1.1','D1.1_2'], 'extra': 'E2.2'},
            {'code': 'PC3', 'desc': '锚固失效+振动', 'deduction': -5, 'triggers': ['C3.2_1','O1.3_1'], 'extra': 'E4'},
            {'code': 'PC4', 'desc': '膨胀土+干湿循环', 'deduction': -7, 'triggers': ['D5.1','E1.2'], 'extra': 'E2.3'},
            {'code': 'PC5', 'desc': '岩溶+爆破(GX-2/4)', 'deduction': -10, 'triggers': ['D5.2'], 'extra': 'E4.1,zone in GX-2/4'},
            {'code': 'PC6', 'desc': '汇水面改变+排水失效', 'deduction': -6, 'triggers': ['E2.1_2','E2.1_3','E2.1_4'], 'extra': '排水失效'},
            {'code': 'PC7', 'desc': '半填半挖+地下水', 'deduction': -5, 'triggers': ['D6.1','O1.7_1','O1.7_2'], 'extra': 'E2.2'}
        ]}
        session.add(SystemConfig(config_type='coupling_rules', config_data=default))
    session.commit()
    session.close()

def get_system_config(config_type):
    session = Session()
    cfg = session.query(SystemConfig).filter_by(config_type=config_type).first()
    session.close()
    return cfg.config_data if cfg else None

def update_system_config(config_type, data):
    session = Session()
    cfg = session.query(SystemConfig).filter_by(config_type=config_type).first()
    if cfg:
        cfg.config_data = data
        session.commit()
        session.close()
        return True
    return False

# ---------- 配置加载 ----------
def ensure_data_dir():
    if not os.path.exists('data'):
        os.makedirs('data')

# 完整缺陷库（含层级信息）
def load_defects_library():
    ensure_data_dir()
    path = os.path.join('data', 'defects_library.json')
    if os.path.exists(path):
        with open(path, 'r', encoding='utf-8') as f:
            return json.load(f)
    default = [
        {"category":"D","code":"D1.1","level2":"地质层符合性","level3":"实际地质与勘察对比","level4":"局部不符","description":"实际地质与勘察对比-局部不符","base":-5},
        {"category":"D","code":"D1.1_2","level2":"地质层符合性","level3":"实际地质与勘察对比","level4":"严重不符","description":"实际地质与勘察对比-严重不符","base":-15},
        {"category":"D","code":"D2.1","level2":"地质水文符合性","level3":"地下水位/渗流场对比","level4":"局部偏差","description":"地下水位/渗流场对比-局部偏差","base":-5},
        {"category":"D","code":"D2.1_2","level2":"地质水文符合性","level3":"地下水位/渗流场对比","level4":"严重偏差","description":"地下水位/渗流场对比-严重偏差","base":-15},
        {"category":"D","code":"D3.1","level2":"防护工程","level3":"增强（过度防护）","level4":"安全系数远高于规范(Fs>1.5)造成浪费","description":"过度防护","base":-2},
        {"category":"D","code":"D3.2","level2":"防护工程","level3":"削弱（防护不足）","level4":"设计安全系数低于规范下限","description":"防护不足","base":-15},
        {"category":"D","code":"D3.3_1","level2":"防护工程","level3":"设置不合理","level4":"过渡防护（小型边坡设大型抗滑桩）","description":"防护过渡","base":-3},
        {"category":"D","code":"D3.3_2","level2":"防护工程","level3":"设置不合理","level4":"防护不足（坡率过陡无锚固）","description":"防护不足","base":-10},
        {"category":"D","code":"D3.3_3","level2":"防护工程","level3":"设置不合理","level4":"防护类型不适合（膨胀土用浆砌片石）","description":"防护类型不当","base":-8},
        {"category":"D","code":"D4.1_1","level2":"排水工程","level3":"尺寸不足","level4":"小于设计流量80%","description":"排水尺寸<80%","base":-5},
        {"category":"D","code":"D4.1_2","level2":"排水工程","level3":"尺寸不足","level4":"小于设计流量50%","description":"排水尺寸<50%","base":-10},
        {"category":"D","code":"D4.2_1","level2":"排水工程","level3":"未设置","level4":"坡顶无截水沟","description":"未设截水沟","base":-8},
        {"category":"D","code":"D4.2_2","level2":"排水工程","level3":"未设置","level4":"坡面无泄水孔","description":"未设泄水孔","base":-8},
        {"category":"D","code":"D4.2_3","level2":"排水工程","level3":"未设置","level4":"坡脚无排水沟","description":"未设坡脚排水沟","base":-8},
        {"category":"D","code":"D4.3_1","level2":"排水工程","level3":"设置不合理","level4":"位置错误（截水沟设在坡面中部）","description":"排水位置错误","base":-6},
        {"category":"D","code":"D4.3_2","level2":"排水工程","level3":"设置不合理","level4":"无反滤层","description":"无反滤层","base":-4},
        {"category":"D","code":"D4.3_3","level2":"排水工程","level3":"设置不合理","level4":"无出口或出口堵塞（设计缺陷）","description":"无出口或堵塞","base":-5},
        {"category":"D","code":"D5.1","level2":"特殊岩土","level3":"膨胀土","level4":"未做保湿/换填/弱膨胀锚固","description":"膨胀土未处理","base":-15},
        {"category":"D","code":"D5.2","level2":"特殊岩土","level3":"岩溶","level4":"未探测溶洞/未处理隐伏溶洞","description":"岩溶未处理","base":-18},
        {"category":"D","code":"D5.3","level2":"特殊岩土","level3":"红黏土/高液限土","level4":"未进行含水率控制或换填","description":"红黏土未控制","base":-8},
        {"category":"D","code":"D6.1","level2":"边坡规模","level3":"坡级增加","level4":"每增加一级（超过设计），累计≤-5","description":"坡级增加每级","base":-1},
        {"category":"C","code":"C1.1","level2":"边坡超挖","level3":"坡率变大","level4":"比设计值每陡0.1（如1:0.5→1:0.4），累计≤-12","description":"坡率变陡每0.1","base":-3},
        {"category":"C","code":"C1.2","level2":"边坡超挖","level3":"局部凹凸","level4":"凹凸深度＞0.5m且面积＞10m²","description":"局部凹凸","base":-4},
        {"category":"C","code":"C2.1","level2":"边坡欠挖","level3":"坡率变小","level4":"导致排水不畅","description":"坡率变小致排水不畅","base":-2},
        {"category":"C","code":"C2.2","level2":"边坡欠挖","level3":"局部凹凸","level4":"凸出岩块可能崩塌","description":"凸出岩块可能崩塌","base":-2},
        {"category":"C","code":"C3.1_1","level2":"防护工程","level3":"支挡失效","level4":"推移（墙顶位移＞5cm）","description":"支挡推移","base":-5},
        {"category":"C","code":"C3.1_2","level2":"防护工程","level3":"支挡失效","level4":"沉降（不均匀＞3cm）","description":"支挡沉降","base":-5},
        {"category":"C","code":"C3.1_3","level2":"防护工程","level3":"支挡失效","level4":"倾斜（倾斜率＞0.5%）","description":"支挡倾斜","base":-6},
        {"category":"C","code":"C3.1_4","level2":"防护工程","level3":"支挡失效","level4":"倾覆（墙趾隆起、墙身后仰）","description":"支挡倾覆","base":-10},
        {"category":"C","code":"C3.1_5","level2":"防护工程","level3":"支挡失效","level4":"破坏（墙体断裂、贯穿裂缝）","description":"支挡破坏","base":-12},
        {"category":"C","code":"C3.1_6","level2":"防护工程","level3":"支挡失效","level4":"开裂（裂缝宽＞2mm，长＞5m）","description":"支挡开裂","base":-4},
        {"category":"C","code":"C3.2_1","level2":"防护工程","level3":"锚固失效","level4":"松弛（预应力损失＞20%）","description":"锚固松弛","base":-4},
        {"category":"C","code":"C3.2_2","level2":"防护工程","level3":"锚固失效","level4":"拉脱（锚头与垫板脱离）","description":"锚固拉脱","base":-8},
        {"category":"C","code":"C3.2_3","level2":"防护工程","level3":"锚固失效","level4":"锈蚀断裂（锚索断裂一根以上）","description":"锚固锈断","base":-10},
        {"category":"C","code":"C3.2_4","level2":"防护工程","level3":"锚固失效","level4":"锚点周边碎裂（混凝土框架开裂剥落）","description":"锚点碎裂","base":-6},
        {"category":"C","code":"C3.3_1","level2":"防护工程","level3":"表层防护失效","level4":"表层冲刷（冲蚀深＞5cm，面积＞5m²）","description":"表层冲刷","base":-3},
        {"category":"C","code":"C3.3_2","level2":"防护工程","level3":"表层防护失效","level4":"表层滑动（局部滑移脱离坡面）","description":"表层滑动","base":-5},
        {"category":"C","code":"C3.3_3","level2":"防护工程","level3":"表层防护失效","level4":"土体裸露（防护层剥落面积＞20%）","description":"土体裸露","base":-4},
        {"category":"C","code":"C3.3_4","level2":"防护工程","level3":"表层防护失效","level4":"开裂渗水（裂缝＞3mm且渗水）","description":"开裂渗水","base":-4},
        {"category":"C","code":"C3.3_5","level2":"防护工程","level3":"表层防护失效","level4":"局部崩塌（小块土体崩落）","description":"局部崩塌","base":-6},
        {"category":"C","code":"C3.4_1","level2":"防护工程","level3":"未按设计施工","level4":"材料不符（强度低一级）","description":"材料不符","base":-6},
        {"category":"C","code":"C3.4_2","level2":"防护工程","level3":"未按设计施工","level4":"位置偏移（抗滑桩偏移＞1m）","description":"位置偏移","base":-4},
        {"category":"C","code":"C3.4_3","level2":"防护工程","level3":"未按设计施工","level4":"数量不足（锚杆少20%）","description":"数量不足","base":-8},
        {"category":"C","code":"C3.4_4","level2":"防护工程","level3":"未按设计施工","level4":"结构形式改变（扶壁→重力式）","description":"结构形式改变","base":-10},
        {"category":"C","code":"C4.1_1","level2":"排水工程","level3":"排水失效","level4":"泄水孔堵塞（堵塞＞30%）","description":"泄水孔堵塞","base":-4},
        {"category":"C","code":"C4.1_2","level2":"排水工程","level3":"排水失效","level4":"渗沟盲沟堵塞（淤积＞20%）","description":"渗沟盲沟堵塞","base":-5},
        {"category":"C","code":"C4.1_3","level2":"排水工程","level3":"排水失效","level4":"无反滤效果（粒料缺失或土工布破损）","description":"无反滤","base":-5},
        {"category":"C","code":"C4.1_4","level2":"排水工程","level3":"排水失效","level4":"排水沟堵塞（堵塞长度＞30%）","description":"排水沟堵塞","base":-6},
        {"category":"C","code":"C4.2_1","level2":"排水工程","level3":"未按设计施工","level4":"位置偏移＞2m","description":"排水位置偏移","base":-5},
        {"category":"C","code":"C4.2_2","level2":"排水工程","level3":"未按设计施工","level4":"范围不足（长度/深度＜80%设计）","description":"排水范围不足","base":-5},
        {"category":"C","code":"C4.2_3","level2":"排水工程","level3":"未按设计施工","level4":"数量不足（泄水孔少20%）","description":"排水数量不足","base":-5},
        {"category":"C","code":"C4.2_4","level2":"排水工程","level3":"未按设计施工","level4":"材料不符（管材/混凝土强度不足）","description":"排水材料不符","base":-5},
        {"category":"C","code":"C4.2_5","level2":"排水工程","level3":"未按设计施工","level4":"结构形式不符（透水要求不满足）","description":"排水结构不符","base":-6},
        {"category":"E","code":"E1.1","level2":"气候变化","level3":"岩层温差裂解","level4":"日温差＞15℃且岩性为泥岩/页岩/炭质岩","description":"岩层温差裂解","base":-2},
        {"category":"E","code":"E1.2","level2":"气候变化","level3":"土层缩胀变形","level4":"膨胀土出现可见裂缝（宽度＞1cm）","description":"膨胀土缩胀","base":-4},
        {"category":"E","code":"E1.3","level2":"气候变化","level3":"温差循环开裂","level4":"混凝土/浆砌体出现明显温差裂缝","description":"温差循环开裂","base":-2},
        {"category":"E","code":"E2.1_1","level2":"水环境变化","level3":"汇水面变化","level4":"排水体系失效（整体排水功能丧失）","description":"排水体系失效","base":-8},
        {"category":"E","code":"E2.1_2","level2":"水环境变化","level3":"汇水面变化","level4":"冲刷变化：汇水面改变（上游汇水面积+30%）","description":"汇水面改变","base":-6},
        {"category":"E","code":"E2.1_3","level2":"水环境变化","level3":"汇水面变化","level4":"冲刷变化：冲槽形成（深度＞0.3m集中冲沟）","description":"冲槽形成","base":-5},
        {"category":"E","code":"E2.1_4","level2":"水环境变化","level3":"汇水面变化","level4":"冲刷变化：结构物支护面掏空（挡墙基础悬空）","description":"结构物掏空","base":-8},
        {"category":"E","code":"E2.1_5","level2":"水环境变化","level3":"区域降雨量超标","level4":"同期雨量超历史（近30年最大）","description":"降雨超历史","base":-8},
        {"category":"E","code":"E2.1_6","level2":"水环境变化","level3":"区域降雨量超标","level4":"最大雨量超设计（小时雨量＞重现期）","description":"降雨超设计","base":-10},
        {"category":"E","code":"E2.1_7","level2":"水环境变化","level3":"区域降雨量超标","level4":"山洪汇水（坡顶临时性洪水流）","description":"山洪汇水","base":-8},
        {"category":"E","code":"E2.1_8","level2":"水环境变化","level3":"区域降雨量超标","level4":"淤塞积水（坡面低洼积水＞0.3m，持续＞3天）","description":"淤塞积水","base":-4},
        {"category":"E","code":"E2.2_1","level2":"水环境变化","level3":"土层软化","level4":"层间粘结弱化→出现推移（坡体向下挤压）","description":"软化推移","base":-5},
        {"category":"E","code":"E2.2_2","level2":"水环境变化","level3":"土层软化","level4":"层间摩擦弱化→出现滑移（整体滑动）","description":"软化滑移","base":-6},
        {"category":"E","code":"E2.2_3","level2":"水环境变化","level3":"土层软化","level4":"层间支撑弱化→出现沉陷（不均匀下沉＞5cm）","description":"软化沉陷","base":-5},
        {"category":"E","code":"E2.2_4","level2":"水环境变化","level3":"土层软化","level4":"层间空隙加大→压缩变形（沉降＞5cm）","description":"软化压缩","base":-4},
        {"category":"E","code":"E2.2_5","level2":"水环境变化","level3":"土层软化","level4":"土层水解软化→软塑或流塑（手捏可塑）","description":"水解软化","base":-6},
        {"category":"E","code":"E2.3_1","level2":"水环境变化","level3":"干湿循环","level4":"湿重变化：整体湿重过大（含水率＞液限）","description":"整体湿重过大","base":-4},
        {"category":"E","code":"E2.3_2","level2":"水环境变化","level3":"干湿循环","level4":"湿重变化：层间湿重不均（含水率差＞10%）","description":"层间湿重不均","base":-3},
        {"category":"E","code":"E2.3_3","level2":"水环境变化","level3":"干湿循环","level4":"土层变形（出现网纹状胀缩裂缝）","description":"胀缩裂缝","base":-5},
        {"category":"E","code":"E2.3_4","level2":"水环境变化","level3":"干湿循环","level4":"表层开裂松散（坡面龟裂，手抓即散）","description":"表层开裂松散","base":-4},
        {"category":"E","code":"E2.4_1","level2":"水环境变化","level3":"地势高差","level4":"动水压力潮汐作用（临河路基水位变动带）","description":"动水潮汐","base":-6},
        {"category":"E","code":"E2.4_2","level2":"水环境变化","level3":"地势高差","level4":"高低位置湿重不平衡（半填挖路基）","description":"湿重不平衡","base":-5},
        {"category":"E","code":"E2.4_3","level2":"水环境变化","level3":"地势高差","level4":"渗流场变化（高大山体透水边坡，坡面溢水）","description":"渗流场变化","base":-6},
        {"category":"E","code":"E2.4_4","level2":"水环境变化","level3":"地势高差","level4":"水头压力差（浸水路堤两侧水位差＞2m）","description":"水头压力差","base":-7},
        {"category":"E","code":"E3.1_1","level2":"堆载变化","level3":"外部堆载","level4":"弃土场（堆高＞3m且距坡顶＜10m）","description":"弃土场堆载","base":-8},
        {"category":"E","code":"E3.1_2","level2":"堆载变化","level3":"外部堆载","level4":"增加结构物（新建电塔/桥墩）","description":"新增结构物","base":-5},
        {"category":"E","code":"E3.1_3","level2":"堆载变化","level3":"外部堆载","level4":"填平区加载（厚度＞2m）","description":"填平区加载","base":-4},
        {"category":"E","code":"E3.2_1","level2":"堆载变化","level3":"坡上负荷变化","level4":"冲刷堆积（坡面堆积物厚度＞0.5m）","description":"冲刷堆积","base":-3},
        {"category":"E","code":"E3.2_2","level2":"堆载变化","level3":"坡上负荷变化","level4":"大平台设置（平台堆载）","description":"大平台堆载","base":-2},
        {"category":"E","code":"E3.3_1","level2":"堆载变化","level3":"坡率变化","level4":"冲刷改变（坡脚冲刷导致坡率变陡，每5°）","description":"冲刷变陡","base":-2},
        {"category":"E","code":"E3.3_2","level2":"堆载变化","level3":"坡率变化","level4":"局部崩塌沉陷（塌陷区面积＞20m²）","description":"局部崩塌沉陷","base":-5},
        {"category":"E","code":"E3.4","level2":"堆载变化","level3":"植被增加","level4":"乔木生长（树高＞3m，可能破坏防护）","description":"乔木生长","base":-3},
        {"category":"E","code":"E4.1","level2":"振动影响","level3":"振能过大失稳","level4":"爆破振动速度＞2cm/s","description":"爆破振动","base":-6},
        {"category":"E","code":"E4.2","level2":"振动影响","level3":"持续振动失稳","level4":"距铁路/重载公路＜50m，且日交通量＞5000辆","description":"持续振动","base":-4},
        {"category":"E","code":"E5.1","level2":"其他因素","level3":"综合因素+时间","level4":"岩土风化严重（手捏可碎，锤击易碎）","description":"岩土风化","base":-3},
        {"category":"E","code":"E5.2","level2":"其他因素","level3":"其他因素","level4":"特殊人为破坏等（由评估人员酌情）","description":"其他因素","base":-2},
        {"category":"O","code":"O1.1_1","level2":"边坡技术状况","level3":"水土流失","level4":"渗水（坡面持续渗水）","description":"渗水","base":-2},
        {"category":"O","code":"O1.1_2","level2":"边坡技术状况","level3":"水土流失","level4":"冲刷（冲沟深度＞0.2m）","description":"冲刷","base":-2},
        {"category":"O","code":"O1.1_3","level2":"边坡技术状况","level3":"水土流失","level4":"溶蚀（岩质边坡出现溶孔、溶槽）","description":"溶蚀","base":-6},
        {"category":"O","code":"O1.2_1","level2":"边坡技术状况","level3":"防护工程（支护变形失效）","level4":"局部失效：沉降（＞2cm）","description":"防护沉降","base":-4},
        {"category":"O","code":"O1.2_2","level2":"边坡技术状况","level3":"防护工程（支护变形失效）","level4":"局部失效：开裂（缝宽＞2mm）","description":"防护开裂","base":-4},
        {"category":"O","code":"O1.2_3","level2":"边坡技术状况","level3":"防护工程（支护变形失效）","level4":"局部失效：推移（＞3cm）","description":"防护推移","base":-5},
        {"category":"O","code":"O1.2_4","level2":"边坡技术状况","level3":"防护工程（支护变形失效）","level4":"局部失效：破损（混凝土剥落＞0.5m²）","description":"防护破损","base":-4},
        {"category":"O","code":"O1.2_5","level2":"边坡技术状况","level3":"防护工程（支护变形失效）","level4":"完全失效（整体滑移、倾覆）","description":"防护完全失效","base":-12},
        {"category":"O","code":"O1.3_1","level2":"边坡技术状况","level3":"锚固体系失效","level4":"局部失效：松弛（预应力损失＞20%）","description":"锚固松弛","base":-4},
        {"category":"O","code":"O1.3_2","level2":"边坡技术状况","level3":"锚固体系失效","level4":"局部失效：松脱（锚头松动）","description":"锚固松脱","base":-4},
        {"category":"O","code":"O1.3_3","level2":"边坡技术状况","level3":"锚固体系失效","level4":"局部失效：开裂（锚墩裂缝）","description":"锚固开裂","base":-3},
        {"category":"O","code":"O1.3_4","level2":"边坡技术状况","level3":"锚固体系失效","level4":"局部失效：推移（锚垫板位移）","description":"锚固推移","base":-4},
        {"category":"O","code":"O1.3_5","level2":"边坡技术状况","level3":"锚固体系失效","level4":"局部失效：破损（锚头锈蚀）","description":"锚固破损","base":-3},
        {"category":"O","code":"O1.3_6","level2":"边坡技术状况","level3":"锚固体系失效","level4":"完全失效（锚索断裂、拔出）","description":"锚固完全失效","base":-12},
        {"category":"O","code":"O1.4_1","level2":"边坡技术状况","level3":"表面防护失效","level4":"局部失效：沉降","description":"表面沉降","base":-3},
        {"category":"O","code":"O1.4_2","level2":"边坡技术状况","level3":"表面防护失效","level4":"局部失效：渗水","description":"表面渗水","base":-3},
        {"category":"O","code":"O1.4_3","level2":"边坡技术状况","level3":"表面防护失效","level4":"局部失效：开裂","description":"表面开裂","base":-3},
        {"category":"O","code":"O1.4_4","level2":"边坡技术状况","level3":"表面防护失效","level4":"局部失效：推移","description":"表面推移","base":-4},
        {"category":"O","code":"O1.4_5","level2":"边坡技术状况","level3":"表面防护失效","level4":"局部失效：破损","description":"表面破损","base":-3},
        {"category":"O","code":"O1.4_6","level2":"边坡技术状况","level3":"表面防护失效","level4":"完全失效（大面积剥落，无防护作用）","description":"表面完全失效","base":-8},
        {"category":"O","code":"O1.5_1","level2":"边坡技术状况","level3":"绿化防护","level4":"稀疏草灌（覆盖率＜50%）","description":"绿化稀疏","base":-1},
        {"category":"O","code":"O1.5_2","level2":"边坡技术状况","level3":"绿化防护","level4":"无植被（完全裸露）","description":"无植被","base":-3},
        {"category":"O","code":"O1.5_3","level2":"边坡技术状况","level3":"绿化防护","level4":"有乔木（树高＞2m，侵入防护）","description":"有乔木","base":-2},
        {"category":"O","code":"O1.6_1","level2":"边坡技术状况","level3":"排水工程失效","level4":"局部失效：沉降（排水沟下沉＞3cm）","description":"排水沉降","base":-3},
        {"category":"O","code":"O1.6_2","level2":"边坡技术状况","level3":"排水工程失效","level4":"局部失效：渗水（沟壁渗漏）","description":"排水渗水","base":-2},
        {"category":"O","code":"O1.6_3","level2":"边坡技术状况","level3":"排水工程失效","level4":"局部失效：开裂（裂缝贯通）","description":"排水开裂","base":-3},
        {"category":"O","code":"O1.6_4","level2":"边坡技术状况","level3":"排水工程失效","level4":"局部失效：推移（沟体位移）","description":"排水推移","base":-4},
        {"category":"O","code":"O1.6_5","level2":"边坡技术状况","level3":"排水工程失效","level4":"局部失效：破损（沟壁崩塌）","description":"排水破损","base":-3},
        {"category":"O","code":"O1.6_6","level2":"边坡技术状况","level3":"排水工程失效","level4":"局部失效：淤积（深度＞断面1/3）","description":"排水淤积","base":-3},
        {"category":"O","code":"O1.6_7","level2":"边坡技术状况","level3":"排水工程失效","level4":"局部失效：堵塞（完全堵塞）","description":"排水堵塞","base":-4},
        {"category":"O","code":"O1.6_8","level2":"边坡技术状况","level3":"排水工程失效","level4":"完全失效（整体断裂、无反滤）","description":"排水完全失效","base":-8},
        {"category":"O","code":"O1.7_1","level2":"边坡技术状况","level3":"边坡规模变化","level4":"坡率大小变化：陡峭（较原设计每陡0.05），累计≤-10","description":"坡率变陡","base":-2},
        {"category":"O","code":"O1.7_2","level2":"边坡技术状况","level3":"边坡规模变化","level4":"分级高度/数量变化：增大（每级增加2m或多一级），累计≤-10","description":"分级增大","base":-2}
    ]
    with open(path, 'w', encoding='utf-8') as f:
        json.dump(default, f, ensure_ascii=False, indent=2)
    return default

def load_default_weights():
    ensure_data_dir()
    path = os.path.join('data', 'default_weights.json')
    if os.path.exists(path):
        with open(path, 'r', encoding='utf-8') as f:
            return json.load(f)
    default = {
        "weight_D":0.20, "weight_C":0.25, "weight_E":0.35, "weight_O":0.20,
        "weight_a1":0.40, "weight_a2":0.30, "weight_a3":0.20, "weight_a4":0.10,
        "zone_factors": {
            "GX-1": {"D1.1":1.3, "D1.1_2":1.3, "E2.1_6":1.3, "E2.4_3":1.3},
            "GX-2": {"D5.2":1.5, "O1.1_3":1.5, "E2.1_6":1.2},
            "GX-3": {"D5.1":1.4, "E1.2":1.4, "E2.3_3":1.4},
            "GX-4": {"D5.2":1.4, "D1.1":1.4, "D1.1_2":1.4},
            "GX-5": {"E2.1_6":1.5, "E2.1_5":1.5, "E2.1_7":1.5},
            "GX-6": {"E4.1":1.3, "E4.2":1.3}
        },
        "coupling_deductions": {"PC1":-8,"PC2":-6,"PC3":-5,"PC4":-7,"PC5":-10,"PC6":-6,"PC7":-5}
    }
    with open(path, 'w', encoding='utf-8') as f:
        json.dump(default, f, ensure_ascii=False, indent=2)
    return default


def init_db():
    # 让 SQLAlchemy 自动创建所有表（如果不存在）
    Base.metadata.create_all(engine)

    session = Session()
    if session.query(SurveyDefect).filter_by(project_id=None).count() == 0:
        defects = load_defects_library()
        for d in defects:
            sd = SurveyDefect(
                project_id=None,
                category=d['category'],
                code=d['code'],
                description=d.get('description', ''),
                level2=d.get('level2', ''),
                level3=d.get('level3', ''),
                level4=d.get('level4', ''),
                base_deduction=d['base'],
                zone_factor=1.0
            )
            session.add(sd)
        session.commit()
    if session.query(WeightConfig).filter_by(is_global=True).count() == 0:
        w = load_default_weights()
        wc = WeightConfig(
            is_global=True,
            weight_D=w['weight_D'], weight_C=w['weight_C'], weight_E=w['weight_E'], weight_O=w['weight_O'],
            weight_a1=w['weight_a1'], weight_a2=w['weight_a2'], weight_a3=w['weight_a3'], weight_a4=w['weight_a4'],
            zone_factors=w['zone_factors'],
            coupling_deductions=w['coupling_deductions']
        )
        session.add(wc)
        session.commit()
    session.close()
    init_system_configs()

# ---------- 计算引擎 ----------
class RiskEngine:
    def __init__(self, project_id):
        self.project_id = project_id
        self.session = Session()
        self.project = self.session.query(Project).filter_by(id=project_id).first()
        if not self.project:
            raise ValueError("Project not found")
        self.wcfg = self.session.query(WeightConfig).filter(
            (WeightConfig.project_id == project_id) | (WeightConfig.is_global == True)
        ).order_by(WeightConfig.is_global.desc()).first()
        if not self.wcfg:
            self.wcfg = WeightConfig(is_global=True)
            self.session.add(self.wcfg)
            self.session.commit()
        self.defects = self.session.query(SurveyDefect).filter_by(project_id=project_id).all()
        self.zone = self.project.zone

    def calc_category_score(self, category):
        total = 0.0
        for d in self.defects:
            if d.category == category and d.checked:
                factor = 1.0
                if self.wcfg.zone_factors and self.zone:
                    zone_factors = self.wcfg.zone_factors.get(self.zone, {})
                    factor = zone_factors.get(d.code, 1.0)
                total += abs(d.base_deduction) * factor
        return max(100 - total, 0)

    def calc_alpha(self):
        L = self.project.current_length or self.project.design_length or 0
        H = self.project.current_height or self.project.design_height or 0
        Lg = '小' if L < 150 else ('中' if L < 250 else '大')
        Hg = '小' if H < 40 else ('中' if H < 55 else '大')
        size_map = {
            ('小','小'):'小', ('小','中'):'小', ('小','大'):'中',
            ('中','小'):'小', ('中','中'):'中', ('中','大'):'大',
            ('大','小'):'中', ('大','中'):'大', ('大','大'):'超大'
        }
        size = size_map.get((Lg, Hg), '中')
        alpha1 = {'小':1.0, '中':0.85, '大':0.70, '超大':0.60}[size]

        tv = self.project.traffic_volume or 0
        if tv < 5000: alpha2 = 1.0
        elif tv < 10000: alpha2 = 0.90
        elif tv < 15000: alpha2 = 0.80
        else: alpha2 = 0.70

        fac_list = self.project.facilities.split(',') if self.project.facilities else ['无']
        fac_map = {'无':1.0, '一般建筑':0.90, '河道':0.90, '沟渠':0.90, '学校':0.80,
                   '医院':0.80, '加油站':0.80, '交叉道路':0.80, '危险品储罐':0.70,
                   '高铁':0.70, '特大桥':0.70, '隧道':0.70}
        alpha3 = min([fac_map.get(f.strip(), 1.0) for f in fac_list]) if fac_list else 1.0

        me = self.project.monitor_emergency or '人工无演练'
        alpha4 = {'有自动化监测 + 应急联动完整（近1年演练）':1.0, '仅有部分（人工、轻量化监测或有预案但无演练或资源缺失）':0.90, '无自动化监测 且 无应急联动（无预案/无演练/无资源）':0.80}.get(me, 0.80)

        w = self.wcfg
        alpha = (w.weight_a1*alpha1 + w.weight_a2*alpha2 + w.weight_a3*alpha3 + w.weight_a4*alpha4) / \
                (w.weight_a1 + w.weight_a2 + w.weight_a3 + w.weight_a4)
        alpha = max(0.5, min(1.0, alpha))
        return alpha1, alpha2, alpha3, alpha4, alpha

    def detect_couplings(self):
        codes = {d.code for d in self.defects if d.checked}
        active = []
        details = []
        coupling_rules = get_system_config('coupling_rules')
        if not coupling_rules:
            return [], []
        for rule in coupling_rules.get('rules', []):
            code = rule['code']
            triggers = rule.get('triggers', [])
            extra = rule.get('extra', '')
            condition_met = False
            reason = ""
            if any(c in codes for c in triggers):
                if extra == 'rainfall_7d>100' and self.project.rainfall_7d and self.project.rainfall_7d > 100:
                    condition_met = True
                    reason = f"触发缺陷 {','.join([c for c in triggers if c in codes])} 且 7日雨量 {self.project.rainfall_7d}mm > 100mm"
                elif extra == 'E2.2' and any(c in codes for c in ['E2.2_1','E2.2_2','E2.2_3','E2.2_4','E2.2_5','E2.4_3']):
                    condition_met = True
                    reason = f"触发缺陷 {','.join([c for c in triggers if c in codes])} 且 存在地下水影响(E2.2)缺陷"
                elif extra == 'E4' and any(c in codes for c in ['E4.1','E4.2']):
                    condition_met = True
                    reason = f"触发缺陷 {','.join([c for c in triggers if c in codes])} 且 存在振动影响(E4)缺陷"
                elif extra == 'E2.3' and any(c in codes for c in ['E2.3_1','E2.3_2','E2.3_3','E2.3_4']):
                    condition_met = True
                    reason = f"触发缺陷 {','.join([c for c in triggers if c in codes])} 且 存在干湿循环(E2.3)缺陷"
                elif extra == 'E4.1,zone in GX-2/4' and 'E4.1' in codes and self.zone in ['GX-2','GX-4']:
                    condition_met = True
                    reason = f"触发缺陷 {','.join([c for c in triggers if c in codes])} 且 分区为 {self.zone} 且存在 E4.1"
                elif extra == '排水失效' and any(c in codes for c in ['D4.2_1','C4.1_4','O1.6_7']):
                    condition_met = True
                    reason = f"触发缺陷 {','.join([c for c in triggers if c in codes])} 且 存在排水失效缺陷"
                elif not extra:
                    condition_met = True
                    reason = f"触发缺陷 {','.join([c for c in triggers if c in codes])}"
            if condition_met:
                active.append(code)
                details.append(f"{code}：{rule['desc']} → 成立。{reason} → 扣 {rule['deduction']} 分")
        return active, details

    def compute(self):
        sD = self.calc_category_score('D')
        sC = self.calc_category_score('C')
        sE = self.calc_category_score('E')
        sO = self.calc_category_score('O')
        w = self.wcfg
        S_base = sD*w.weight_D + sC*w.weight_C + sE*w.weight_E + sO*w.weight_O

        coupling_rules = get_system_config('coupling_rules')
        coupling_map = {}
        for rule in coupling_rules.get('rules', []):
            coupling_map[rule['code']] = rule['deduction']
        active, details = self.detect_couplings()
        coupling_total = sum(coupling_map.get(c,0) for c in active)
        S_base_corrected = max(S_base + coupling_total, 0)

        a1, a2, a3, a4, alpha = self.calc_alpha()
        S_final = S_base_corrected * alpha

        grade_cfg = get_system_config('grade_thresholds')
        grade = '未定义'
        for t in grade_cfg.get('thresholds', []):
            if t['min'] <= S_final <= t['max']:
                grade = t['level']
                break
        if grade == '未定义' and S_final >= 80:
            grade = '低风险'
        elif grade == '未定义' and S_final >= 60:
            grade = '中风险'
        elif grade == '未定义' and S_final >= 40:
            grade = '高风险'
        elif grade == '未定义':
            grade = '极高风险'

        final_grade = grade
        if self.project.fs_input is not None:
            fs = self.project.fs_input
            design_fs = 1.25
            if fs < 1.0:
                final_grade = '极高风险'
            elif fs < design_fs:
                order = ['低风险','中风险','高风险','极高风险']
                idx = order.index(grade) if grade in order else 0
                if idx < 3:
                    final_grade = order[idx+1]
            elif fs >= design_fs * 1.1:
                order = ['低风险','中风险','高风险','极高风险']
                idx = order.index(grade) if grade in order else 0
                if idx > 0:
                    final_grade = order[idx-1]

        measures_cfg = get_system_config('response_measures')
        response = measures_cfg.get(final_grade, '请配置响应措施')

        result = Result(
            project_id=self.project_id,
            score_D=sD, score_C=sC, score_E=sE, score_O=sO,
            S_base=S_base,
            coupling_total=coupling_total,
            S_base_corrected=S_base_corrected,
            alpha1=a1, alpha2=a2, alpha3=a3, alpha4=a4, alpha=alpha,
            S_final=S_final,
            initial_grade=grade,
            final_grade=final_grade,
            response_measures=response
        )
        self.session.add(result)
        self.session.commit()
        self.session.close()
        return result, details

# ---------- 导入导出 ----------
def generate_import_template():
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        info = {
            '项目名称': [''], '路线名称': [''], '桩号': [''],
            '左右侧': ['左侧'], '边坡类型': ['路堑'], '评估分区': [''],
            '设计高度(m)': [''], '当前高度(m)': [''], '设计长度(m)': [''],
            '当前长度(m)': [''], '设计坡率': [''], '实际坡率': [''],
            '设计分级数量': [''], '实际分级数量': [''],
            '日均交通量(pcu)': [''], '周边重要设施': [''],
            '监测应急能力': [''], '运营年限': [''],
            '前3天雨量(mm)': [''], '前7天雨量(mm)': [''],
            '力学复核Fs': ['']
        }
        pd.DataFrame({k: pd.Series(v) for k,v in info.items()}).to_excel(writer, sheet_name='基本信息', index=False)
        session = Session()
        defects = session.query(SurveyDefect).filter_by(project_id=None).all()
        session.close()
        data = [{'类别':d.category,'编号':d.code,'二级指标':d.level2,'三级指标':d.level3,'四级指标':d.level4,
                 '描述':d.description,'是否存在(填1或0)':0,'基础扣分':d.base_deduction} for d in defects]
        pd.DataFrame(data).to_excel(writer, sheet_name='缺陷勾选', index=False)
    output.seek(0)
    return output

def import_from_excel(file_bytes, project_id):
    df_info = pd.read_excel(file_bytes, sheet_name='基本信息', header=0)
    df_defects = pd.read_excel(file_bytes, sheet_name='缺陷勾选', header=0)
    info = df_info.iloc[0].to_dict()
    for k,v in info.items():
        if pd.isna(v):
            info[k] = None
    session = Session()
    project = session.query(Project).filter_by(id=project_id).first()
    if not project:
        session.close()
        return False, "项目不存在"
    mapping = {
        '项目名称':'name','路线名称':'route','桩号':'pile_no',
        '左右侧':'side','边坡类型':'type','评估分区':'zone',
        '设计高度(m)':'design_height','当前高度(m)':'current_height',
        '设计长度(m)':'design_length','当前长度(m)':'current_length',
        '设计坡率':'design_slope','实际坡率':'actual_slope',
        '设计分级数量':'grade_count','实际分级数量':'platform_count',
        '日均交通量(pcu)':'traffic_volume','周边重要设施':'facilities',
        '监测应急能力':'monitor_emergency','运营年限':'operating_years',
        '前3天雨量(mm)':'rainfall_3d','前7天雨量(mm)':'rainfall_7d',
        '力学复核Fs':'fs_input'
    }
    for excel_col, db_col in mapping.items():
        if excel_col in info and info[excel_col] is not None:
            setattr(project, db_col, info[excel_col])
    for idx, row in df_defects.iterrows():
        code = row['编号']
        exists = row['是否存在(填1或0)']
        if pd.isna(exists):
            continue
        defect = session.query(SurveyDefect).filter_by(project_id=project_id, code=code).first()
        if defect:
            defect.checked = bool(int(exists))
        else:
            lib_def = session.query(SurveyDefect).filter_by(project_id=None, code=code).first()
            if lib_def:
                new_def = SurveyDefect(
                    project_id=project_id,
                    category=lib_def.category,
                    code=lib_def.code,
                    description=lib_def.description,
                    level2=lib_def.level2,
                    level3=lib_def.level3,
                    level4=lib_def.level4,
                    base_deduction=lib_def.base_deduction,
                    checked=bool(int(exists)),
                    zone_factor=1.0
                )
                session.add(new_def)
    session.commit()
    session.close()
    return True, "导入成功"

def export_project_data(project_id):
    session = Session()
    project = session.query(Project).filter_by(id=project_id).first()
    defects = session.query(SurveyDefect).filter_by(project_id=project_id).all()
    session.close()
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        info = {
            '项目名称': project.name, '路线名称': project.route, '桩号': project.pile_no,
            '左右侧': project.side, '边坡类型': project.type, '评估分区': project.zone,
            '设计高度(m)': project.design_height, '当前高度(m)': project.current_height,
            '设计长度(m)': project.design_length, '当前长度(m)': project.current_length,
            '设计坡率': project.design_slope, '实际坡率': project.actual_slope,
            '设计分级数量': project.grade_count, '实际分级数量': project.platform_count,
            '日均交通量(pcu)': project.traffic_volume, '周边重要设施': project.facilities,
            '监测应急能力': project.monitor_emergency, '运营年限': project.operating_years,
            '前3天雨量(mm)': project.rainfall_3d, '前7天雨量(mm)': project.rainfall_7d,
            '力学复核Fs': project.fs_input
        }
        pd.DataFrame([info]).to_excel(writer, sheet_name='基本信息', index=False)
        data = [{'类别':d.category,'编号':d.code,'二级指标':d.level2,'三级指标':d.level3,'四级指标':d.level4,
                 '描述':d.description,'是否存在':1 if d.checked else 0,'基础扣分':d.base_deduction,
                 '分区系数':d.zone_factor} for d in defects]
        pd.DataFrame(data).to_excel(writer, sheet_name='缺陷勾选', index=False)
    output.seek(0)
    return output

def export_results(project_id):
    session = Session()
    project = session.query(Project).filter_by(id=project_id).first()
    result = session.query(Result).filter_by(project_id=project_id).order_by(Result.id.desc()).first()
    defects = session.query(SurveyDefect).filter_by(project_id=project_id).all()
    wcfg = session.query(WeightConfig).filter(
        (WeightConfig.project_id == project_id) | (WeightConfig.is_global == True)
    ).order_by(WeightConfig.is_global.desc()).first()
    session.close()
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        info = {'项目名称':project.name,'桩号':project.pile_no,'分区':project.zone,
                '设计高度':project.design_height,'当前高度':project.current_height,
                '设计长度':project.design_length,'当前长度':project.current_length,
                '交通量':project.traffic_volume,'周边设施':project.facilities,
                '监测应急':project.monitor_emergency}
        pd.DataFrame([info]).to_excel(writer, sheet_name='项目信息', index=False)
        if result:
            res = {'得分D':result.score_D,'得分C':result.score_C,'得分E':result.score_E,'得分O':result.score_O,
                   'S_base':result.S_base,'耦合扣分':result.coupling_total,
                   'S_base_corrected':result.S_base_corrected,
                   'α1':result.alpha1,'α2':result.alpha2,'α3':result.alpha3,'α4':result.alpha4,
                   'α':result.alpha,'S_final':result.S_final,
                   '初判等级':result.initial_grade,'最终等级':result.final_grade,
                   '响应措施':result.response_measures}
            pd.DataFrame([res]).to_excel(writer, sheet_name='评估结果', index=False)
        defect_list = [{'类别':d.category,'编号':d.code,'描述':d.description,
                        '是否勾选':d.checked,'基础扣分':d.base_deduction,
                        '分区系数':d.zone_factor,
                        '实际扣分':d.base_deduction*d.zone_factor if d.checked else 0} for d in defects]
        pd.DataFrame(defect_list).to_excel(writer, sheet_name='缺陷明细', index=False)
        w = {'D权重':wcfg.weight_D,'C权重':wcfg.weight_C,'E权重':wcfg.weight_E,'O权重':wcfg.weight_O,
             'a1权重':wcfg.weight_a1,'a2权重':wcfg.weight_a2,'a3权重':wcfg.weight_a3,'a4权重':wcfg.weight_a4}
        pd.DataFrame([w]).to_excel(writer, sheet_name='权重配置', index=False)
    output.seek(0)
    return output

def generate_word_report(project_id):
    session = Session()
    project = session.query(Project).filter_by(id=project_id).first()
    result = session.query(Result).filter_by(project_id=project_id).order_by(Result.id.desc()).first()
    defects = session.query(SurveyDefect).filter_by(project_id=project_id).all()
    wcfg = session.query(WeightConfig).filter(
        (WeightConfig.project_id == project_id) | (WeightConfig.is_global == True)
    ).order_by(WeightConfig.is_global.desc()).first()
    grade_cfg = get_system_config('grade_thresholds')
    safety_cfg = get_system_config('safety_factors')
    engine = RiskEngine(project_id)
    _, details = engine.compute()
    session.close()
    doc = Document()
    title = doc.add_heading('广西高速公路边坡安全评估报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'报告编号：{datetime.now().strftime("%Y%m%d%H%M%S")}')
    doc.add_paragraph(f'评估日期：{datetime.now().strftime("%Y年%m月%d日")}')
    doc.add_paragraph()

    doc.add_heading('一、项目概要', level=1)
    doc.add_paragraph(f'项目名称：{project.name}')
    doc.add_paragraph(f'路线：{project.route}，桩号：{project.pile_no}')
    doc.add_paragraph(f'边坡类型：{project.type}，分区：{project.zone}')
    doc.add_paragraph(f'设计高度：{project.design_height}m，当前高度：{project.current_height}m')
    doc.add_paragraph(f'设计长度：{project.design_length}m，当前长度：{project.current_length}m')
    doc.add_paragraph(f'运营年限：{project.operating_years}年')
    doc.add_paragraph(f'日均交通量：{project.traffic_volume} pcu')
    doc.add_paragraph(f'周边重要设施：{project.facilities}')
    doc.add_paragraph(f'监测应急能力：{project.monitor_emergency}')

    if result:
        doc.add_heading('二、评估结论', level=1)
        doc.add_paragraph(f'最终安全指数 S_final = {result.S_final:.2f}')
        doc.add_paragraph(f'风险等级：{result.final_grade}')
        for t in grade_cfg.get('thresholds', []):
            if t['level'] == result.final_grade:
                doc.add_paragraph(f'预警色：{t.get("color", "未配置")}')
                break
        doc.add_paragraph(f'响应措施：{result.response_measures}')

        doc.add_heading('三、评分明细', level=1)
        data_rows = [
            ('设计符合性 D', result.score_D),
            ('施工质量 C', result.score_C),
            ('环境变化 E', result.score_E),
            ('运营现状 O', result.score_O),
            ('加权基础分 S_base', result.S_base)
        ]
        table = doc.add_table(rows=len(data_rows) + 1, cols=2)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text = '一级指标'
        hdr[1].text = '得分'
        for i, (name, val) in enumerate(data_rows, start=1):
            row = table.rows[i].cells
            row[0].text = name
            row[1].text = f'{val:.2f}'

        doc.add_heading('四、详细计算过程', level=1)
        doc.add_paragraph(f'加权基础分 S_base = {result.score_D:.2f}×{wcfg.weight_D:.2f} + {result.score_C:.2f}×{wcfg.weight_C:.2f} + {result.score_E:.2f}×{wcfg.weight_E:.2f} + {result.score_O:.2f}×{wcfg.weight_O:.2f} = {result.S_base:.2f}')
        doc.add_paragraph('耦合效应修正：')
        if details:
            for d in details:
                doc.add_paragraph(f'  {d}')
        else:
            doc.add_paragraph('  未触发任何耦合效应。')
        doc.add_paragraph(f'耦合扣分合计：{result.coupling_total:.2f}')
        doc.add_paragraph(f'修正后 S_base_corrected = {result.S_base_corrected:.2f}')
        doc.add_paragraph('后果折减系数 α 计算：')
        doc.add_paragraph(f'  规模等级：根据长度{project.current_length}m、高度{project.current_height}m 判定，α1={result.alpha1}')
        doc.add_paragraph(f'  交通流量：日均{project.traffic_volume}辆，α2={result.alpha2}')
        doc.add_paragraph(f'  周边设施：{project.facilities}，α3={result.alpha3}')
        doc.add_paragraph(f'  监测应急：{project.monitor_emergency}，α4={result.alpha4}')
        doc.add_paragraph(f'  α = {wcfg.weight_a1:.2f}×{result.alpha1} + {wcfg.weight_a2:.2f}×{result.alpha2} + {wcfg.weight_a3:.2f}×{result.alpha3} + {wcfg.weight_a4:.2f}×{result.alpha4} = {result.alpha:.3f}')
        doc.add_paragraph(f'最终 S_final = {result.S_base_corrected:.2f} × {result.alpha:.3f} = {result.S_final:.2f}')
        
        doc.add_heading('五、等级判定标准', level=1)
        table2 = doc.add_table(rows=len(grade_cfg.get('thresholds', [])) + 1, cols=3)
        table2.style = 'Light Grid Accent 1'
        hdr2 = table2.rows[0].cells
        hdr2[0].text = 'S_final 范围'
        hdr2[1].text = '风险等级'
        hdr2[2].text = '预警色'
        for i, t in enumerate(grade_cfg.get('thresholds', []), start=1):
            row = table2.rows[i].cells
            row[0].text = f"{t['min']} ~ {t['max']}"
            row[1].text = t['level']
            row[2].text = t.get('color', '')
        doc.add_paragraph(f'当前 S_final = {result.S_final:.2f}，落在 {result.final_grade} 区间。')
        
        if project.fs_input is not None:
            doc.add_heading('六、力学复核与等级调整', level=1)
            doc.add_paragraph(f'输入力学复核安全系数 Fs = {project.fs_input}')
            for k, v in safety_cfg.items():
                doc.add_paragraph(f'{k}：{v}')
            fs = project.fs_input
            design_fs = 1.25
            if fs < 1.0:
                doc.add_paragraph('Fs < 1.0，直接定为极高风险。')
            elif fs < design_fs:
                doc.add_paragraph(f'1.0 ≤ Fs ({fs}) < 设计值 ({design_fs})，等级上调一级。')
            elif fs >= design_fs * 1.1:
                doc.add_paragraph(f'Fs ({fs}) ≥ 设计值×1.1 ({design_fs*1.1:.2f})，等级下调一级。')
            else:
                doc.add_paragraph(f'Fs ({fs}) 满足设计值要求，等级不做调整。')
            doc.add_paragraph(f'最终等级：{result.final_grade}')

    checked = [d for d in defects if d.checked]
    doc.add_heading('七、现场调查缺陷记录', level=1)
    if checked:
        table3 = doc.add_table(rows=1+len(checked), cols=6)
        table3.style = 'Light Grid Accent 1'
        hdr3 = table3.rows[0].cells
        hdr3[0].text = '类别'
        hdr3[1].text = '编号'
        hdr3[2].text = '四级指标'
        hdr3[3].text = '描述'
        hdr3[4].text = '扣分'
        hdr3[5].text = '系数'
        for i, d in enumerate(checked, start=1):
            row = table3.rows[i].cells
            row[0].text = d.category
            row[1].text = d.code
            row[2].text = d.level4 or d.description
            row[3].text = d.description
            row[4].text = str(d.base_deduction * d.zone_factor)
            row[5].text = str(d.zone_factor)
    else:
        doc.add_paragraph('未发现缺陷。')

    doc.add_heading('八、权重配置', level=1)
    doc.add_paragraph(f'一级指标权重：D={wcfg.weight_D}, C={wcfg.weight_C}, E={wcfg.weight_E}, O={wcfg.weight_O}')
    doc.add_paragraph(f'后果折减权重：α1={wcfg.weight_a1}, α2={wcfg.weight_a2}, α3={wcfg.weight_a3}, α4={wcfg.weight_a4}')
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# ---------- 主页面 ----------
def main():
    st.set_page_config(page_title="边坡风险评估系统", layout="wide")
    st.sidebar.title("导航")
    page = st.sidebar.radio("跳转", ["项目管理", "数据录入", "参数配置", "结果与报告"])

    if page == "项目管理":
        show_project_management()
    elif page == "数据录入":
        show_data_entry()
    elif page == "参数配置":
        show_parameter_config()
    elif page == "结果与报告":
        show_result_report()

# ---------- 项目管理 ----------
def show_project_management():
    st.title("项目管理")
    session = Session()
    with st.expander("新建项目"):
        name = st.text_input("项目名称")
        if st.button("创建"):
            if name:
                try:
                    proj = Project(name=name)
                    session.add(proj)
                    session.commit()
                    lib_defs = session.query(SurveyDefect).filter_by(project_id=None).all()
                    for d in lib_defs:
                        new_d = SurveyDefect(
                            project_id=proj.id,
                            category=d.category,
                            code=d.code,
                            description=d.description,
                            level2=d.level2,
                            level3=d.level3,
                            level4=d.level4,
                            base_deduction=d.base_deduction,
                            checked=False,
                            zone_factor=1.0
                        )
                        session.add(new_d)
                    global_w = session.query(WeightConfig).filter_by(is_global=True).first()
                    if global_w:
                        wc = WeightConfig(
                            project_id=proj.id,
                            weight_D=global_w.weight_D,
                            weight_C=global_w.weight_C,
                            weight_E=global_w.weight_E,
                            weight_O=global_w.weight_O,
                            weight_a1=global_w.weight_a1,
                            weight_a2=global_w.weight_a2,
                            weight_a3=global_w.weight_a3,
                            weight_a4=global_w.weight_a4,
                            zone_factors=global_w.zone_factors,
                            coupling_deductions=global_w.coupling_deductions,
                            is_global=False
                        )
                        session.add(wc)
                    session.commit()
                    st.success(f"项目 '{name}' 创建成功！请点击左侧“数据录入”继续。")
                except Exception as e:
                    st.error(f"创建失败：{str(e)}")
            else:
                st.warning("请输入项目名称")
    
    st.subheader("已有项目")
    projects = session.query(Project).all()
    if not projects:
        st.info("暂无项目")
    else:
        data = []
        # 原来的循环查询
        # for p in projects:
            # result = session.query(Result).filter_by(project_id=p.id).order_by(Result.id.desc()).first()
        # 改为一次性查询（需要调整逻辑）
        from sqlalchemy import func, desc
        subq = session.query(Result.project_id, func.max(Result.id).label('max_id')).group_by(
            Result.project_id).subquery()
        results = session.query(Result).join(subq, (Result.id == subq.c.max_id) & (
                    Result.project_id == subq.c.project_id)).all()
        result_map = {r.project_id: r for r in results}
        # 然后在循环中从 result_map 获取
            grade = result.final_grade if result else "未评估"
            data.append({"ID": p.id, "名称": p.name, "桩号": p.pile_no, "等级": grade, "更新时间": p.updated_at.strftime("%Y-%m-%d")})
        df = pd.DataFrame(data)
        selected_ids = st.multiselect(
            "选择要删除的项目（可多选）",
            options=df['ID'].tolist(),
            format_func=lambda x: f"ID {x} - {df[df['ID']==x]['名称'].iloc[0]}"
        )
        col1, col2, col3 = st.columns([1,1,1])
        with col1:
            if st.button("删除选中项目"):
                if selected_ids:
                    try:
                        for pid in selected_ids:
                            session.query(SurveyDefect).filter_by(project_id=pid).delete()
                            session.query(Result).filter_by(project_id=pid).delete()
                            session.query(WeightConfig).filter_by(project_id=pid).delete()
                            session.query(Project).filter_by(id=pid).delete()
                        session.commit()
                        st.success(f"成功删除 {len(selected_ids)} 个项目")
                        st.rerun()
                    except Exception as e:
                        st.error(f"删除失败：{str(e)}")
                else:
                    st.warning("请至少选择一个项目")
        with col2:
            if st.button("刷新列表"):
                st.rerun()
        with col3:
            st.info("在复选框中勾选多个ID即可批量操作")
        
        st.subheader("加载项目")
        sel_id = st.number_input("输入项目ID加载", min_value=1, step=1)
        if st.button("加载该项目"):
            if sel_id in df['ID'].tolist():
                st.session_state['project_id'] = sel_id
                st.success(f"已加载项目 ID={sel_id}")
            else:
                st.error("项目ID不存在")
        st.dataframe(df, use_container_width=True)
    
    st.subheader("批量导入项目")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("下载批量导入模板"):
            template = generate_import_template()
            st.download_button("点击下载", data=template, file_name="批量导入模板.xlsx", 
                              mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    with col2:
        uploaded = st.file_uploader("上传批量项目Excel", type=["xlsx"])
        if uploaded and st.button("导入批量项目"):
            try:
                df_batch = pd.read_excel(uploaded, sheet_name='基本信息', header=0)
                success_count = 0
                for idx, row in df_batch.iterrows():
                    if pd.isna(row['项目名称']):
                        continue
                    proj = Project(name=row['项目名称'])
                    for col in ['路线名称','桩号','左右侧','边坡类型','评估分区',
                                '设计高度(m)','当前高度(m)','设计长度(m)','当前长度(m)',
                                '日均交通量(pcu)','周边重要设施','监测应急能力','运营年限']:
                        if col in df_batch.columns and not pd.isna(row[col]):
                            mapping = {
                                '路线名称':'route','桩号':'pile_no','左右侧':'side',
                                '边坡类型':'type','评估分区':'zone',
                                '设计高度(m)':'design_height','当前高度(m)':'current_height',
                                '设计长度(m)':'design_length','当前长度(m)':'current_length',
                                '日均交通量(pcu)':'traffic_volume','周边重要设施':'facilities',
                                '监测应急能力':'monitor_emergency','运营年限':'operating_years'
                            }
                            if col in mapping:
                                setattr(proj, mapping[col], row[col])
                    session.add(proj)
                    session.flush()
                    lib_defs = session.query(SurveyDefect).filter_by(project_id=None).all()
                    for d in lib_defs:
                        new_d = SurveyDefect(
                            project_id=proj.id,
                            category=d.category,
                            code=d.code,
                            description=d.description,
                            level2=d.level2,
                            level3=d.level3,
                            level4=d.level4,
                            base_deduction=d.base_deduction,
                            checked=False,
                            zone_factor=1.0
                        )
                        session.add(new_d)
                    global_w = session.query(WeightConfig).filter_by(is_global=True).first()
                    if global_w:
                        wc = WeightConfig(
                            project_id=proj.id,
                            weight_D=global_w.weight_D,
                            weight_C=global_w.weight_C,
                            weight_E=global_w.weight_E,
                            weight_O=global_w.weight_O,
                            weight_a1=global_w.weight_a1,
                            weight_a2=global_w.weight_a2,
                            weight_a3=global_w.weight_a3,
                            weight_a4=global_w.weight_a4,
                            zone_factors=global_w.zone_factors,
                            coupling_deductions=global_w.coupling_deductions,
                            is_global=False
                        )
                        session.add(wc)
                    success_count += 1
                session.commit()
                st.success(f"成功导入 {success_count} 个项目")
                st.rerun()
            except Exception as e:
                st.error(f"导入失败：{str(e)}")
    session.close()

# ---------- 数据录入 ----------
def show_data_entry():
    st.title("数据录入")
    if 'project_id' not in st.session_state:
        st.warning("请先在项目管理中加载一个项目")
        return
    pid = st.session_state['project_id']
    session = Session()
    project = session.query(Project).filter_by(id=pid).first()
    if not project:
        st.error("项目不存在")
        return
    st.subheader(f"项目：{project.name} (ID={pid})")
    
    with st.form("basic_info"):
        col1, col2 = st.columns(2)
        with col1:
            project.route = st.text_input("路线名称", project.route)
            project.pile_no = st.text_input("桩号", project.pile_no)
            project.side = st.selectbox("左右侧", ["左侧","右侧"], index=0 if project.side=="左侧" else 1)
            project.type = st.selectbox("边坡类型", ["路堑","路堤"], index=0 if project.type=="路堑" else 1)
            project.zone = st.selectbox("评估分区", ["GX-1","GX-2","GX-3","GX-4","GX-5","GX-6"],
                                         index=["GX-1","GX-2","GX-3","GX-4","GX-5","GX-6"].index(project.zone) if project.zone in ["GX-1","GX-2","GX-3","GX-4","GX-5","GX-6"] else 0)
            project.design_height = st.number_input("设计高度(m)", value=project.design_height or 0.0)
            project.current_height = st.number_input("当前高度(m)", value=project.current_height or 0.0)
            project.design_length = st.number_input("设计长度(m)", value=project.design_length or 0.0)
            project.current_length = st.number_input("当前长度(m)", value=project.current_length or 0.0)
        with col2:
            project.design_slope = st.text_input("设计坡率", project.design_slope)
            project.actual_slope = st.text_input("实际坡率", project.actual_slope)
            project.grade_count = st.number_input("设计分级数量", value=project.grade_count or 0, step=1)
            project.platform_count = st.number_input("实际分级数量", value=project.platform_count or 0, step=1)
            project.traffic_volume = st.number_input("日均交通量(pcu)", value=project.traffic_volume or 0, step=100)
            facilities_options = ['无','一般建筑','河道','沟渠','学校','医院','加油站','交叉道路','危险品储罐','高铁','特大桥','隧道']
            current_fac = project.facilities.split(',') if project.facilities else ['无']
            selected_fac = st.multiselect("周边重要设施（可多选）", facilities_options, default=current_fac)
            project.facilities = ','.join(selected_fac) if selected_fac else '无'
            monitor_opts = ['自动化', '人工有演练', '人工无演练']
            current_mon = project.monitor_emergency if project.monitor_emergency in monitor_opts else '人工无演练'
            project.monitor_emergency = st.selectbox("监测应急能力", monitor_opts, index=monitor_opts.index(current_mon))
            project.operating_years = st.number_input("运营年限", value=project.operating_years or 0, step=1)
            project.rainfall_3d = st.number_input("前3天累计雨量(mm)", value=project.rainfall_3d or 0.0)
            project.rainfall_7d = st.number_input("前7天累计雨量(mm)", value=project.rainfall_7d or 0.0)
            project.fs_input = st.number_input("力学复核Fs(可选)", value=project.fs_input or 0.0, step=0.01)
        if st.form_submit_button("保存基本信息"):
            session.commit()
            st.success("基本信息已保存")
    
    st.subheader("缺陷勾选")
    defects = session.query(SurveyDefect).filter_by(project_id=pid).all()
    if not defects:
        st.info("请先初始化缺陷库")
        return
    cat_order = {"D":"设计符合性", "C":"施工质量", "E":"环境变化", "O":"运营现状"}
    for cat in ["D","C","E","O"]:
        with st.expander(f"{cat_order[cat]} (扣分项)"):
            cat_defs = [d for d in defects if d.category == cat]
            l2_groups = {}
            for d in cat_defs:
                key = d.level2 or "未分类"
                if key not in l2_groups:
                    l2_groups[key] = []
                l2_groups[key].append(d)
            for l2, items in l2_groups.items():
                st.markdown(f"**{l2}**")
                l3_groups = {}
                for d in items:
                    key = d.level3 or "未分类"
                    if key not in l3_groups:
                        l3_groups[key] = []
                    l3_groups[key].append(d)
                for l3, items3 in l3_groups.items():
                    st.markdown(f"&nbsp;&nbsp;*{l3}*")
                    cols = st.columns(2)
                    for i, d in enumerate(items3):
                        col = cols[i % 2]
                        wcfg = session.query(WeightConfig).filter(
                            (WeightConfig.project_id == pid) | (WeightConfig.is_global == True)
                        ).order_by(WeightConfig.is_global.desc()).first()
                        zone_factors = wcfg.zone_factors or {}
                        factor = zone_factors.get(project.zone, {}).get(d.code, 1.0)
                        d.zone_factor = factor
                        label = f"{d.code} - {d.level4 or d.description} (扣{abs(d.base_deduction)*factor:.1f})"
                        checked = col.checkbox(label, value=d.checked, key=f"def_{d.id}")
                        if checked != d.checked:
                            d.checked = checked
                            session.commit()
                st.markdown("---")
    
    st.subheader("数据导入导出")
    col1, col2, col3 = st.columns(3)
    with col1:
        if st.button("下载导入模板"):
            template = generate_import_template()
            st.download_button("点击下载", data=template, file_name="导入模板.xlsx", 
                              mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    with col2:
        uploaded = st.file_uploader("上传已填写的Excel", type=["xlsx"])
        if uploaded and st.button("导入数据"):
            success, msg = import_from_excel(uploaded, pid)
            if success:
                st.success(msg)
                st.rerun()
            else:
                st.error(msg)
    with col3:
        if st.button("导出当前项目数据"):
            data = export_project_data(pid)
            st.download_button("下载项目数据", data=data, file_name=f"项目_{project.name}_数据.xlsx",
                              mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    session.close()

# ---------- 参数配置 ----------
def show_parameter_config():
    st.title("参数配置")
    if 'project_id' not in st.session_state:
        st.warning("请先加载一个项目")
        return
    pid = st.session_state['project_id']
    session = Session()
    # 获取全局权重配置
    wcfg = session.query(WeightConfig).filter_by(is_global=True).first()
    if not wcfg:
        st.error("全局权重配置不存在，请重新初始化数据库。")
        session.close()
        return

    # 初始化 session_state 中的显示数据
    if 'zone_factors_display' not in st.session_state:
        st.session_state['zone_factors_display'] = wcfg.zone_factors or {}

    # ----- 一级指标权重 -----
    st.subheader("一级指标权重 (总和应为1)")
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        wD = st.number_input("设计 D", value=wcfg.weight_D, min_value=0.0, max_value=1.0, step=0.01)
    with col2:
        wC = st.number_input("施工 C", value=wcfg.weight_C, min_value=0.0, max_value=1.0, step=0.01)
    with col3:
        wE = st.number_input("环境 E", value=wcfg.weight_E, min_value=0.0, max_value=1.0, step=0.01)
    with col4:
        wO = st.number_input("运营 O", value=wcfg.weight_O, min_value=0.0, max_value=1.0, step=0.01)
    total = wD + wC + wE + wO
    st.write(f"当前总和：{total:.2f}")
    if abs(total - 1.0) > 0.01:
        st.warning("权重之和不等于1，请调整")
    else:
        if st.button("保存一级权重"):
            wcfg.weight_D, wcfg.weight_C, wcfg.weight_E, wcfg.weight_O = wD, wC, wE, wO
            session.commit()
            st.success("一级权重已更新")

    # ----- 后果折减因子权重 -----
    st.subheader("后果折减因子权重")
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        wa1 = st.number_input("α1 规模", value=wcfg.weight_a1, min_value=0.0, max_value=1.0, step=0.05)
    with col2:
        wa2 = st.number_input("α2 交通", value=wcfg.weight_a2, min_value=0.0, max_value=1.0, step=0.05)
    with col3:
        wa3 = st.number_input("α3 设施", value=wcfg.weight_a3, min_value=0.0, max_value=1.0, step=0.05)
    with col4:
        wa4 = st.number_input("α4 监测", value=wcfg.weight_a4, min_value=0.0, max_value=1.0, step=0.05)
    total_a = wa1 + wa2 + wa3 + wa4
    st.write(f"当前总和：{total_a:.2f}")
    if abs(total_a - 1.0) > 0.01:
        st.warning("权重之和不等于1，请调整")
    else:
        if st.button("保存后果权重"):
            wcfg.weight_a1, wcfg.weight_a2, wcfg.weight_a3, wcfg.weight_a4 = wa1, wa2, wa3, wa4
            session.commit()
            st.success("后果权重已更新")

    # ----- 广西六大评估分区调整系数表 -----
    st.subheader("广西六大评估分区调整系数表")
    st.info("在此配置各分区下特定缺陷编码的系数，计算时将自动乘以对应扣分。")
    zones = ["GX-1", "GX-2", "GX-3", "GX-4", "GX-5", "GX-6"]

    # 从 session_state 获取显示数据
    zone_factors_display = st.session_state['zone_factors_display']

    # ---------- 修改或新增系数 ----------
    st.write("**修改或新增系数**")
    col1, col2, col3 = st.columns(3)
    with col1:
        edit_zone = st.selectbox("选择分区", zones)
    with col2:
        # 从数据库获取缺陷编码，保证下拉列表非空
        all_defects = session.query(SurveyDefect).filter_by(project_id=None).all()
        defect_codes = [d.code for d in all_defects] if all_defects else ["D1.1", "C1.1", "E1.1", "O1.1_1"]
        edit_code = st.selectbox("选择缺陷编码", defect_codes, key='edit_code')
    with col3:
        edit_factor = st.number_input("系数值", value=1.0, step=0.1, key='edit_factor')

    # 显示成功消息
    if 'add_success' in st.session_state:
        st.success(st.session_state['add_success'])
        del st.session_state['add_success']

    if st.button("添加/更新该编码系数", key="add_factor_btn"):
        # 更新 session_state 中的字典
        if edit_zone not in zone_factors_display:
            zone_factors_display[edit_zone] = {}
        zone_factors_display[edit_zone][edit_code] = edit_factor
        # 保存到数据库
        wcfg.zone_factors = zone_factors_display
        session.commit()
        # 更新 session_state
        st.session_state['zone_factors_display'] = zone_factors_display
        st.session_state['add_success'] = f"✅ 已更新 {edit_zone} 下 {edit_code} 系数为 {edit_factor}"
        st.rerun()

    # 显示当前配置（从 session_state 读取）
    st.write("**当前配置**（仅显示已配置的编码）")
    # 使用 Expander 展示清单
    zone_factors_display = st.session_state.get('zone_factors_display', {})
    for zone in ["GX-1", "GX-2", "GX-3", "GX-4", "GX-5", "GX-6"]:
        with st.expander(f"{zone} 当前系数"):
            if zone in zone_factors_display and zone_factors_display[zone]:
                for code, factor in zone_factors_display[zone].items():
                    st.write(f"{code}: {factor}")
            else:
                st.write("无特殊系数")
    # ---------- 删除已有系数 ----------
    with st.expander("删除已有系数"):
        del_zone = st.selectbox("选择分区", zones, key="del_zone_select")
        # 获取当前分区下的编码列表
        current_factors = st.session_state.get('zone_factors_display', {})
        if del_zone in current_factors and current_factors[del_zone]:
            del_code = st.selectbox("选择要删除的编码", list(current_factors[del_zone].keys()), key="del_code_select")
            if st.button("删除该系数", key="del_factor_btn"):
                # 从 session_state 中删除
                del current_factors[del_zone][del_code]
                if not current_factors[del_zone]:
                    del current_factors[del_zone]
                # 更新数据库
                wcfg.zone_factors = current_factors
                session.commit()
                # 更新 session_state
                st.session_state['zone_factors_display'] = current_factors
                st.success(f"已删除 {del_zone} 下的 {del_code}")
                st.rerun()
        else:
            st.info("该分区暂无配置的系数")

    # ----- 等级阈值配置 -----
    st.subheader("等级划分阈值配置")
    grade_cfg = get_system_config('grade_thresholds')
    if grade_cfg:
        new_thresholds = []
        for i, t in enumerate(grade_cfg.get('thresholds', [])):
            col1, col2, col3, col4 = st.columns([1, 1, 1, 0.5])
            with col1:
                min_val = st.number_input(f"最小值{i + 1}", value=t['min'], key=f"grade_min_{i}")
            with col2:
                max_val = st.number_input(f"最大值{i + 1}", value=t['max'], key=f"grade_max_{i}")
            with col3:
                level_val = st.text_input(f"等级名称{i + 1}", value=t['level'], key=f"grade_level_{i}")
            with col4:
                color_val = st.text_input(f"颜色{i + 1}", value=t.get('color', ''), key=f"grade_color_{i}")
            new_thresholds.append({'min': min_val, 'max': max_val, 'level': level_val, 'color': color_val})
        if st.button("保存等级阈值"):
            update_system_config('grade_thresholds', {'thresholds': new_thresholds})
            st.success("等级阈值已更新")

    # ----- 响应措施配置 -----
    st.subheader("响应措施配置")
    measures_cfg = get_system_config('response_measures')
    if measures_cfg:
        new_measures = {}
        for level, measure in measures_cfg.items():
            if isinstance(measure, str):
                measure = [measure] if measure.strip() else []
            measure_text = "\n".join(measure) if isinstance(measure, list) else measure
            new_text = st.text_area(
                f"{level} 响应措施（每行一条）",
                value=measure_text,
                height=150,
                key=f"measure_{level}"
            )
            new_measures[level] = [line.strip() for line in new_text.splitlines() if line.strip()]
        if st.button("保存响应措施"):
            update_system_config('response_measures', new_measures)
            st.success("响应措施已更新")
    else:
        st.info("暂无响应措施配置，请先初始化。")

    # ----- 安全系数阈值配置 -----
    st.subheader("安全系数阈值配置")
    safety_cfg = get_system_config('safety_factors')
    if safety_cfg:
        new_safety = {}
        for name, value in safety_cfg.items():
            new_safety[name] = st.number_input(f"{name} 安全系数", value=value, step=0.01, key=f"safety_{name}")
        if st.button("保存安全系数阈值"):
            update_system_config('safety_factors', new_safety)
            st.success("安全系数阈值已更新")

    # ----- 耦合效应规则配置 -----
    st.subheader("耦合效应规则配置")
    coupling_cfg = get_system_config('coupling_rules')
    if coupling_cfg:
        new_rules = []
        for i, rule in enumerate(coupling_cfg.get('rules', [])):
            with st.expander(f"规则 {i + 1}: {rule['code']} - {rule['desc']}"):
                code = st.text_input("代码", value=rule['code'], key=f"cpl_code_{i}")
                desc = st.text_input("描述", value=rule['desc'], key=f"cpl_desc_{i}")
                deduction = st.number_input("扣分值", value=rule['deduction'], step=1, key=f"cpl_ded_{i}")
                triggers = st.text_input("触发缺陷编码（逗号分隔）", value=','.join(rule.get('triggers', [])),
                                         key=f"cpl_trig_{i}")
                extra = st.text_input("额外条件", value=rule.get('extra', ''), key=f"cpl_extra_{i}")
                new_rules.append({
                    'code': code,
                    'desc': desc,
                    'deduction': deduction,
                    'triggers': [x.strip() for x in triggers.split(',') if x.strip()],
                    'extra': extra
                })
        st.write("---")
        st.write("**新增自定义耦合规则**")
        with st.expander("添加新规则"):
            new_code = st.text_input("新规则代码（如：PC8）", key="new_cpl_code")
            new_desc = st.text_input("新规则描述", key="new_cpl_desc")
            new_ded = st.number_input("新规则扣分值", value=-5, step=1, key="new_cpl_ded")
            new_trig = st.text_input("触发缺陷编码（逗号分隔）", key="new_cpl_trig")
            new_extra = st.text_input("额外条件", key="new_cpl_extra")
            if st.button("添加新规则"):
                if new_code and new_desc:
                    new_rules.append({
                        'code': new_code,
                        'desc': new_desc,
                        'deduction': new_ded,
                        'triggers': [x.strip() for x in new_trig.split(',') if x.strip()],
                        'extra': new_extra
                    })
                    st.success(f"规则 {new_code} 已添加到列表，请点击下方保存按钮生效")
        if st.button("保存耦合规则"):
            update_system_config('coupling_rules', {'rules': new_rules})
            st.success("耦合规则已更新")

    session.close()

# ---------- 结果与报告 ----------
def show_result_report():
    st.title("结果与报告")
    if 'project_id' not in st.session_state:
        st.warning("请先加载一个项目")
        return
    pid = st.session_state['project_id']
    session = Session()
    project = session.query(Project).filter_by(id=pid).first()
    if not project:
        st.error("项目不存在")
        return
    st.subheader(f"项目：{project.name}")
    
    if st.button("开始评估"):
        try:
            engine = RiskEngine(pid)
            result, details = engine.compute()
            st.session_state['coupling_details'] = details
            st.success("评估完成！请查看下方结果。")
        except Exception as e:
            st.error(f"评估失败：{str(e)}")
    
    result = session.query(Result).filter_by(project_id=pid).order_by(Result.id.desc()).first()
    if not result:
        st.info("尚未进行评估，请点击上方按钮")
        session.close()
        return
    
    coupling_details = st.session_state.get('coupling_details', [])
    
    col1, col2 = st.columns(2)
    with col1:
        st.metric("最终安全指数 S_final", f"{result.S_final:.2f}")
        st.metric("风险等级", result.final_grade)
    with col2:
        st.metric("初判等级", result.initial_grade)
        st.metric("耦合扣分", f"{result.coupling_total:.1f}")
    
    wcfg = session.query(WeightConfig).filter(
        (WeightConfig.project_id == pid) | (WeightConfig.is_global == True)
    ).order_by(WeightConfig.is_global.desc()).first()
    
    st.subheader("一级指标得分表")
    df_score = pd.DataFrame({
        "指标": ["设计符合性 D", "施工质量 C", "环境变化 E", "运营现状 O"],
        "得分": [result.score_D, result.score_C, result.score_E, result.score_O],
        "权重": [wcfg.weight_D if wcfg else 0.20, wcfg.weight_C if wcfg else 0.25, 
                 wcfg.weight_E if wcfg else 0.35, wcfg.weight_O if wcfg else 0.20],
        "加权得分": [result.score_D * (wcfg.weight_D if wcfg else 0.20),
                    result.score_C * (wcfg.weight_C if wcfg else 0.25),
                    result.score_E * (wcfg.weight_E if wcfg else 0.35),
                    result.score_O * (wcfg.weight_O if wcfg else 0.20)]
    })
    st.dataframe(df_score, use_container_width=True)
    
    st.subheader("详细计算过程")
    st.write(f"**加权基础分 S_base** = {result.score_D:.2f}×{wcfg.weight_D:.2f} + {result.score_C:.2f}×{wcfg.weight_C:.2f} + {result.score_E:.2f}×{wcfg.weight_E:.2f} + {result.score_O:.2f}×{wcfg.weight_O:.2f} = **{result.S_base:.2f}**")
    
    st.write("**耦合效应修正：**")
    if coupling_details:
        for d in coupling_details:
            st.write(f"- {d}")
    else:
        st.write("未触发任何耦合效应。")
    st.write(f"耦合扣分合计：**{result.coupling_total:.2f}**")
    st.write(f"**修正后 S_base_corrected** = {result.S_base_corrected:.2f}")
    
    st.write("**后果折减系数 α 计算：**")
    st.write(f"- 规模等级：α1 = {result.alpha1}")
    st.write(f"- 交通流量：α2 = {result.alpha2}")
    st.write(f"- 周边设施：α3 = {result.alpha3}")
    st.write(f"- 监测应急：α4 = {result.alpha4}")
    st.write(f"α = {wcfg.weight_a1:.2f}×{result.alpha1} + {wcfg.weight_a2:.2f}×{result.alpha2} + {wcfg.weight_a3:.2f}×{result.alpha3} + {wcfg.weight_a4:.2f}×{result.alpha4} = **{result.alpha:.3f}**")
    st.write(f"**最终 S_final** = {result.S_base_corrected:.2f} × {result.alpha:.3f} = **{result.S_final:.2f}**")
    
    st.subheader("等级判定过程")
    grade_cfg = get_system_config('grade_thresholds')
    st.write("依据等级划分标准：")
    grade_table = []
    for t in grade_cfg.get('thresholds', []):
        grade_table.append([f"{t['min']} ~ {t['max']}", t['level'], t.get('color', '')])
    st.table(pd.DataFrame(grade_table, columns=["S_final 范围", "风险等级", "预警色"]))
    st.write(f"当前 S_final = **{result.S_final:.2f}**，判定为 **{result.final_grade}**")
    
    st.subheader("响应措施")
    st.info(result.response_measures)
    # ---------- 最终结论与措施 ----------
    st.subheader("最终结论与措施")
    # 获取该等级对应的措施列表
    measures_cfg = get_system_config('response_measures')
    level_measures = measures_cfg.get(result.final_grade, []) if measures_cfg else []
    if isinstance(level_measures, str):
        level_measures = [level_measures]
    if not level_measures:
        level_measures = ['请配置响应措施']

    # 获取缺陷列表（用于显示致灾因子）
    defects = session.query(SurveyDefect).filter_by(project_id=pid).all()

    col1, col2 = st.columns(2)
    with col1:
        st.markdown(f"**最终等级**：{result.final_grade}")
        st.markdown(
            f"**安全指数**：{result.S_final:.1f}（初判{result.initial_grade}，力学复核{'上调' if result.final_grade != result.initial_grade else '未调整'}）")
    with col2:
        checked_defects = [d for d in defects if d.checked]
        main_factors = []
        for d in checked_defects[:5]:
            desc = d.level4 or d.description
            if desc:
                main_factors.append(desc)
        st.markdown(f"**主要致灾因子**：{'、'.join(main_factors) if main_factors else '未识别'}")

    st.markdown("**应急响应措施：**")
    for idx, item in enumerate(level_measures, 1):
        st.markdown(f"{idx}. {item}")


    col1, col2, col3 = st.columns(3)
    with col1:
        if st.button("导出评估报告(Word)"):
            try:
                word_file = generate_word_report(pid)
                st.download_button("下载报告", data=word_file, file_name="评估报告.docx", 
                                  mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            except Exception as e:
                st.error(f"导出失败：{str(e)}")
    with col2:
        if st.button("导出结果(Excel)"):
            excel_file = export_results(pid)
            st.download_button("下载Excel", data=excel_file, file_name="评估结果.xlsx", 
                              mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    with col3:
        with st.expander("更新力学复核值"):
            fs = st.number_input("输入Fs", value=0.0, step=0.01)
            if st.button("确认更新"):
                try:
                    project.fs_input = fs
                    session.commit()
                    st.success("已更新，请重新点击“开始评估”")
                except Exception as e:
                    st.error(f"更新失败：{str(e)}")
    session.close()

if __name__ == "__main__":
    # 先创建所有表（如果不存在）
    Base.metadata.create_all(engine)
    init_db()
    main()